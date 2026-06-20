"""
resume_extractor.py
───────────────────
Server-side text extraction for resume files.

Supported formats:
  • .pdf  — pdfplumber (layout-aware; falls back to pypdf if needed)
  • .docx — python-docx
  • .doc  — doc2docx conversion → python-docx  (requires LibreOffice)
  • .txt  — plain UTF-8 read (with encoding fallback)
  • .rtf  — striprtf library

Install dependencies:
  pip install pdfplumber pypdf python-docx striprtf
  # Optional for scanned PDFs (OCR):
  pip install pytesseract pdf2image
"""

import io
import os
import re
import subprocess
import tempfile


# ─────────────────────────────────────────────────────────────
# PUBLIC API
# ─────────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> tuple[str, str | None]:
    """
    Extract plain text from a resume file.

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename:   Original filename (used to determine format).

    Returns:
        (text, error)
        - text  : Extracted string (empty string on failure)
        - error : Human-readable error message, or None on success
    """
    ext = _get_ext(filename)

    extractors = {
        '.pdf':  _extract_pdf,
        '.docx': _extract_docx,
        '.doc':  _extract_doc,
        '.txt':  _extract_txt,
        '.rtf':  _extract_rtf,
    }

    fn = extractors.get(ext)
    if fn is None:
        supported = ', '.join(sorted(extractors.keys()))
        return '', f"Unsupported file type '{ext}'. Supported formats: {supported}"

    try:
        text = fn(file_bytes)
        text = _clean(text)
        if not text:
            return '', 'No readable text found in the file. If this is a scanned PDF, please copy-paste the text manually.'
        return text, None
    except Exception as e:
        return '', f"Could not extract text from {ext} file: {str(e)}"


# ─────────────────────────────────────────────────────────────
# EXTRACTORS
# ─────────────────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes) -> str:
    """
    Uses pdfplumber for best layout-aware text extraction.
    Falls back to pypdf if pdfplumber is unavailable or fails.
    Falls back to OCR as a last resort if tesseract is installed.
    """
    # ── Attempt 1: pdfplumber (best for text-based PDFs) ──
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if page_text:
                    pages.append(page_text)
            text = '\n\n'.join(pages)
            if text.strip():
                return text
    except ImportError:
        pass  # pdfplumber not installed — try pypdf
    except Exception:
        pass  # corrupted or unusual PDF — try pypdf

    # ── Attempt 2: pypdf fallback ──
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages.append(t)
        text = '\n\n'.join(pages)
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception:
        pass

    # ── Attempt 3: OCR for scanned PDFs ──
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
        images = convert_from_bytes(file_bytes, dpi=200)
        pages = [pytesseract.image_to_string(img) for img in images]
        return '\n\n'.join(pages)
    except ImportError:
        raise RuntimeError(
            "PDF appears to be a scanned image and OCR libraries are not installed. "
            "Please install pdfplumber and pytesseract, or copy-paste your resume text manually."
        )


def _extract_docx(file_bytes: bytes) -> str:
    """Extracts text from .docx using python-docx, preserving paragraph structure."""
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Also extract text from tables (skills tables, education tables etc.)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                parts.append('  |  '.join(row_cells))

    return '\n'.join(parts)


def _extract_doc(file_bytes: bytes) -> str:
    """
    Extracts text from legacy .doc files.
    Converts to .docx via LibreOffice, then uses python-docx.
    Requires: LibreOffice installed (libreoffice --headless).
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path  = os.path.join(tmpdir, 'resume.doc')
        docx_path = os.path.join(tmpdir, 'resume.docx')

        with open(doc_path, 'wb') as f:
            f.write(file_bytes)

        try:
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'docx',
                 '--outdir', tmpdir, doc_path],
                capture_output=True, timeout=30
            )
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr.decode()}")
        except FileNotFoundError:
            raise RuntimeError(
                ".doc format requires LibreOffice to be installed. "
                "Please save your resume as .docx or .pdf and re-upload."
            )

        if not os.path.exists(docx_path):
            raise RuntimeError("LibreOffice conversion produced no output file.")

        with open(docx_path, 'rb') as f:
            return _extract_docx(f.read())


def _extract_txt(file_bytes: bytes) -> str:
    """Reads plain text with automatic encoding detection."""
    # Try UTF-8 first, then common fallbacks
    for encoding in ('utf-8', 'utf-8-sig', 'latin-1', 'cp1252'):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    # Last resort: decode with replacement characters
    return file_bytes.decode('utf-8', errors='replace')


def _extract_rtf(file_bytes: bytes) -> str:
    """Strips RTF formatting using the striprtf library."""
    try:
        from striprtf.striprtf import rtf_to_text
        rtf_string = file_bytes.decode('latin-1', errors='replace')
        return rtf_to_text(rtf_string)
    except ImportError:
        raise RuntimeError(
            "RTF support requires the 'striprtf' library. "
            "Run: pip install striprtf\n"
            "Or save your resume as .pdf or .docx and re-upload."
        )


# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def _get_ext(filename: str) -> str:
    """Returns the lowercased file extension including the dot."""
    _, ext = os.path.splitext(filename.lower())
    return ext


def _clean(text: str) -> str:
    """
    Normalises extracted text:
    - Collapses 3+ consecutive blank lines to 2
    - Strips leading/trailing whitespace per line
    - Removes null bytes and other control characters
    """
    if not text:
        return ''

    # Remove null bytes and control chars (except newline, tab, carriage return)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    # Normalise Windows line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Strip trailing whitespace from each line
    lines = [line.rstrip() for line in text.split('\n')]

    # Collapse runs of 3+ blank lines to 2
    cleaned = []
    blank_run = 0
    for line in lines:
        if line == '':
            blank_run += 1
            if blank_run <= 2:
                cleaned.append(line)
        else:
            blank_run = 0
            cleaned.append(line)

    return '\n'.join(cleaned).strip()